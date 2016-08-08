# -*- coding: utf-8 -*-
"""
Created on Thu Jul 21 10:30:35 2016
work for the tradedetail with new style
function
1 read the file  as moudle attribution
2 make a main table,index is fund_account
3 make other tables, one person netvalue(day),one person profolio(day-trade)
one perosn profolio(trans)
4  option attribation from wind
5  plot tool
6  pdf show
7  excel show 

@author: yangxs
"""

from datetime import datetime,timedelta
import pyfolio as pf
import pandas as pd
import matplotlib
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
from WindPy import w
import numpy as np
import os
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.oxml.ns import qn
import sys
#rare data 
cdf=pd.DataFrame() #allacct
tdf=pd.DataFrame() #tradedetail
CT=pd.DataFrame()  #custom table
R=0.0   #无风险收益率
dir_url=u'D:/temp/testOP/'
matplotlib.use('Agg')

option_dict={(0L, 'C') : 'call_option_buyer',
(0L, 'P') : 'put_option_buyer',
(1L, 'C') : 'call_option_seller',
(1L, 'P') : 'put_option_seller',
(2L, 'C') : 'covered_call',
(2L, 'P') : 'no_covered_put'
}

def readRaredata(url):
    # use / instead \   
    # read the rare data
    df=pd.read_csv(url)
    #print df    
    print url    
    return df

def filterData(tdf):
    #filter the tdf,because of close date =0
    #fill the close date=0    
    #tdf_f=pd.DataFrame(tdf[tdf.DROP_DATE!=0])
    #return tdf_f  
    tdf.loc[tdf.DROP_DATE==0,'DROP_DATE']=20991231

def dealwithDate(cdf,tdf):
    #deal with date
    cdf[u'date']=pd.to_datetime(cdf.INIT_DATE.apply(str))
    tdf[u'date']=pd.to_datetime(tdf.INIT_DATE.apply(str))
    tdf[u'opendate']=pd.to_datetime(tdf.OPEN_DATE.apply(str))
    tdf[u'closedate']=pd.to_datetime(tdf.DROP_DATE.apply(str))
    tdf[u'dura_day']=tdf.closedate-tdf.opendate
    tdf[u'PNL']=tdf.INCOME_BALANCE.divide(tdf.OPEN_BALANCE.abs())

def calCustomtransaction(CT,tdf):
    """
    所有交易按笔数统计：买入认购、买入认沽、卖出认购、卖出认沽、备对开仓的数量及占比（请注意平仓的不要考虑）。
    所有交易按张数统计：上述类型交易数量及占比
    按所有交易笔数/交易张数分别统计持仓时间处于日内交易，日间但不超过1周的，超过一周的交易张数及其占比

    对持有买入交易/卖出交易/买入认购/买入认沽/卖出认购/卖出认沽/备对开仓/日内仓/非日内仓分别进行下列统计：
    平均持仓时间：
    每笔交易平均盈亏金额：
    每笔交易平均盈亏比率：
    按交易笔数统计胜率：
    按笔数统计盈亏比率：
    所有交易日中持有或者进行此类交易的天数占比：
    此类交易占所有交易的笔数占比：

    行权笔数:
    行权张数:
    被行权指派笔数:
    被行权指派张数:
    
    """
    #开仓笔数和开仓次数
    CT=pd.DataFrame(tdf.groupby(by='FUND_ACCOUNT').OPEN_AMOUNT.sum())
    CT[u'open_num']=pd.DataFrame(tdf.groupby(by='FUND_ACCOUNT').OPEN_AMOUNT.count())
    k3=tdf.groupby(by=['FUND_ACCOUNT','OPTION_TYPE','OPTHOLD_TYPE']).OPEN_AMOUNT.sum()
    k4=k3.unstack().unstack()
    new_columns=['open_amount_'+option_dict[i]  for i in k4.columns.get_values()]
    k4.columns=new_columns
    CT=CT.join(k4)
    
    k7=tdf.groupby(by=['FUND_ACCOUNT','OPTION_TYPE','OPTHOLD_TYPE']).OPEN_AMOUNT.count()
    k8=k7.unstack().unstack()
    new_columns=['open_num_'+option_dict[i]  for i in k8.columns.get_values()]
    k8.columns=new_columns
    CT=CT.join(k8)    
    
    #开仓时间
    k9=tdf.groupby(by=['FUND_ACCOUNT','dura_day']).OPEN_AMOUNT.sum()
    k10=k9.unstack()
    CT[u'trade_in_one_day_amount']=k10[timedelta(0)]
    for i in range(1,8):
        k11=k10[timedelta(i)]
    CT[u'trade_in_week_amount']=k11
    CT[u'trade_more_than_week_amount']=k10.sum(axis=1)-k11-k10[timedelta(0)]
    
    k9=tdf.groupby(by=['FUND_ACCOUNT','dura_day']).OPEN_AMOUNT.count()
    k10=k9.unstack()
    CT[u'trade_in_one_day_num']=k10[timedelta(0)]
    for i in range(1,8):
        k11=k10[timedelta(i)]
    CT[u'trade_in_week_num']=k11
    CT[u'trade_more_than_week_num']=k10.sum(axis=1)-k11-k10[timedelta(0)]
    
    #日内仓，非日内仓统计，dura_time,profit_per_num,PNL,win_ratio，trade_ratio
    CT[u'trade_in_one_day_holddays']=tdf[tdf.dura_day==timedelta(0)]\
    .groupby(by=['FUND_ACCOUNT']).HOLD_DAYS.mean()
    CT[u'trade_more_one_day_holddays']=tdf[tdf.dura_day>timedelta(0)]\
    .groupby(by=['FUND_ACCOUNT']).HOLD_DAYS.mean()    
    CT[u'trade_in_one_day_profit_per_num']=tdf[tdf.dura_day==timedelta(0)]\
    .groupby(by=['FUND_ACCOUNT']).INCOME_BALANCE.mean()
    CT[u'trade_more_one_day_profit_per_num']=tdf[tdf.dura_day>timedelta(0)]\
    .groupby(by=['FUND_ACCOUNT']).INCOME_BALANCE.mean()
    CT[u'trade_in_one_day_pnl']=tdf[tdf.dura_day==timedelta(0)]\
    .groupby(by=['FUND_ACCOUNT']).PNL.mean()
    CT[u'trade_more_one_day_pnl']=tdf[tdf.dura_day>timedelta(0)].\
    groupby(by=['FUND_ACCOUNT']).PNL.mean()
    CT[u'trade_in_one_day_trade_ratio']=tdf[tdf.dura_day==timedelta(0)].\
    groupby(by=['FUND_ACCOUNT']).OPEN_AMOUNT.count().divide(CT.open_num)
    CT[u'trade_more_one_day_trade_ratio']=tdf[tdf.dura_day>timedelta(0)].\
    groupby(by=['FUND_ACCOUNT']).OPEN_AMOUNT.count().divide(CT.open_num)
    CT[u'trade_in_one_day_win_ratio']=tdf[tdf.dura_day==timedelta(0)]\
    [tdf.INCOME_BALANCE>0].groupby(by=['FUND_ACCOUNT']).OPEN_AMOUNT\
    .count()/tdf[tdf.dura_day==timedelta(0)]\
    .groupby(by=['FUND_ACCOUNT']).OPEN_AMOUNT.count()
    CT[u'trade_more_one_day_win_ratio']=tdf[tdf.dura_day>timedelta(0)]\
    [tdf.INCOME_BALANCE>0].groupby(by=['FUND_ACCOUNT']).OPEN_AMOUNT\
    .count()/tdf[tdf.dura_day>timedelta(0)]\
    .groupby(by=['FUND_ACCOUNT']).OPEN_AMOUNT.count()  
    
    #买入认购/买入认沽/卖出认购/卖出认沽/备对开仓的计算
    k21=tdf.groupby(by=['FUND_ACCOUNT','OPTION_TYPE','OPTHOLD_TYPE'])\
    .HOLD_DAYS.mean()   
    k22=k21.unstack().unstack()
    new_columns=['hold_days_'+option_dict[i]  for i in k22.columns.get_values()]
    k22.columns=new_columns
    CT=CT.join(k22)
   
    k23=tdf.groupby(by=['FUND_ACCOUNT','OPTION_TYPE','OPTHOLD_TYPE'])\
    .INCOME_BALANCE.mean()   
    k24=k23.unstack().unstack()
    new_columns=['income_balance_'+option_dict[i]  for i in k24.columns.get_values()]
    k24.columns=new_columns
    CT=CT.join(k24)
    
    k25=tdf.groupby(by=['FUND_ACCOUNT','OPTION_TYPE','OPTHOLD_TYPE'])\
    .PNL.mean()   
    k26=k25.unstack().unstack()
    new_columns=['PNL_'+option_dict[i]  for i in k26.columns.get_values()]
    k26.columns=new_columns
    CT=CT.join(k26)
    
    k27=tdf.groupby(by=['FUND_ACCOUNT','OPTION_TYPE','OPTHOLD_TYPE'])\
    .OPEN_AMOUNT.count().divide(CT.open_num)   
    k28=k27.unstack().unstack()
    new_columns=['trade_ratio_'+option_dict[i]  for i in k28.columns.get_values()]
    k28.columns=new_columns
    CT=CT.join(k28)
   
    k29=tdf[tdf.INCOME_BALANCE>0].groupby(by=['FUND_ACCOUNT','OPTION_TYPE','OPTHOLD_TYPE'])\
    .OPEN_AMOUNT.count()/k7   
    k30=k29.unstack().unstack()
    new_columns=['win_ratio_'+option_dict[i]  for i in k30.columns.get_values()]
    k30.columns=new_columns
    CT=CT.join(k30)   
    
    CT.fillna(0,inplace=True)
    return CT
            
def getPostion(tdf,acct):
    """
    get position by acct 
    calc delta,vega.....
    """
    #tdf[tdf.FUND_ACCOUNT==9008000020][(tdf.opendate<=datetime(2015,4,11))& (tdf.closedate>=datetime(2015,4,11))]
    return
    
def calCustomNV(CT,cdf,tdf):
    """
    每笔交易平均盈亏金额:
    每笔交易平均盈亏比率:
        平均每笔交易持仓时间：
        每天净值收益率平均值:
        净值波动率:
        Sharp ratio
        区间的Beta值
        最大回撤：
        按交易笔数统计胜率：
        按交易笔数统计盈亏比率：
        按天统计胜率：
        按天统计盈亏比率（盈利的天数盈利百分比平均值/亏损天数亏损百分比平均值）
    """    
    
    # cdf,tdf has diffent rows,because some customs didnt trade    
    # std,    
    #CT[u'nv_std']=cdf.groupby(by='FUND_ACCOUNT').REAL_OPT_NET.std()
    #k59=cdf.groupby(by='FUND_ACCOUNT').REAL_OPT_NET.std()
    k59=cdf.groupby(by='FUND_ACCOUNT').REAL_OPT_NET\
    .apply(lambda x :x.pct_change().std())    
    k59.rename('nv_std',inplace=True)
    CT=CT.join(pd.DataFrame(k59),how='outer')
    CT[u'annual_nv_std']=CT.nv_std.mul(245**0.5) 
    CT[u'total_profit']=cdf.groupby(by='FUND_ACCOUNT').INCOME_BALANCE.sum()
    CT[u'opendate']=pd.to_datetime(cdf.groupby(by='FUND_ACCOUNT')\
    .OPEN_DATE.first().apply(str))
    CT[u'first_trade_day']=tdf.groupby(by='FUND_ACCOUNT').opendate.first()
    CT[u'last_trade_day']=tdf.groupby(by='FUND_ACCOUNT').closedate.last()
    CT['last_trade_day']=CT['last_trade_day']\
    .apply(lambda x : x if x!=datetime(2099,12,31) else datetime(2016,6,16))
    CT[u'last_nv']=cdf.groupby(by='FUND_ACCOUNT').REAL_OPT_NET.last()
    CT[u'total_return']=CT.last_nv-1
    k57=tdf.groupby(by=['FUND_ACCOUNT','opendate']).OPEN_AMOUNT.count()
    CT[u'days_of_open']=k57.count(level='FUND_ACCOUNT')
    #interval  std,think about begin and end of trade    
    CT[u'interval_nv_std']=cdf.groupby('FUND_ACCOUNT').apply(f,CT)
    CT[u'annual_interval_nv_std']=CT.interval_nv_std.mul(245**0.5)
    # return of year,get n
    CT[u'interval_days']=CT.last_trade_day-CT.first_trade_day+timedelta(1)
    CT[u'interval_trading_days']=cdf.groupby('FUND_ACCOUNT')\
    .apply(f2,CT)
    CT[u'return_of_year_by_interval']=CT.last_nv**(245/CT.interval_trading_days)-1
    CT[u'sharp_ratio']=(CT.return_of_year_by_interval-R)/CT.annual_interval_nv_std
    
    #CT.fillna(0,inplace=True)
    return CT

def f(x,CT):
    #calc the std for the trading day
    acct=long(x['FUND_ACCOUNT'].head(1).get_values())
    #print x[x>datetime(2016,6,1)]    
    #return x[(x>=CT['first_trade_day'])&(x<=CT['last_trade_day'])]
    """    
    return x[(x.date>=CT.loc[acct].first_trade_day)\
    &(x.date<=CT.loc[acct].last_trade_day)].REAL_OPT_NET\
    .apply(lambda x :x.pct_change().std())      
    """ 
    y=x[(x.date>=CT.loc[acct].first_trade_day)\
    &(x.date<=CT.loc[acct].last_trade_day)].REAL_OPT_NET.pct_change().std()
    #print acct
    return y
     
def f2(x,CT):
    # count the trading day
    acct=long(x['FUND_ACCOUNT'].head(1).get_values())
    y=x[(x.date>=CT.loc[acct].first_trade_day)\
    &(x.date<=CT.loc[acct].last_trade_day)].date.count()
    #print acct
    return y    

def draw1(acct,CT,cdf,tdf):
    #初步绘图
    #isinteractive(False)
    url=dir_url+str(acct)
    os.makedirs(url)
    tt=cdf[cdf.FUND_ACCOUNT==acct]
    #fig1=plt.figure()
    #ax1=fig.add_subplot(1,1,1)
    #net_value
    plot1=tt.plot(x=['date'],y=['REAL_OPT_NET'])
    fig1=plot1.get_figure()
    fig1.savefig(url+'/'+u'净值')
    #持仓希腊字母
    plot2=tt.plot(x=['date'],y=['DELTA'])
    fig2=plot2.get_figure()
    fig2.savefig(url+'/'+u'DELTA')
    plot3=tt.plot(x=['date'],y=['GAMMA'])
    fig3=plot3.get_figure()
    fig3.savefig(url+'/'+u'GAMMA')
    plot4=tt.plot(x=['date'],y=['VEGA'])
    fig4=plot4.get_figure()
    fig4.savefig(url+'/'+u'VEGA')   
    plot5=tt.plot(x=['date'],y=['THETA'])
    fig5=plot5.get_figure()
    fig5.savefig(url+'/'+u'THETA') 
    #交易盈亏
    plot6=tt.plot(x=['date'],y=['TRADE_INCOME'])
    fig6=plot6.get_figure()
    fig6.savefig(url+'/'+u'交易盈亏') 
    #风险度
    plot7=tt.plot(x=['date'],y=['RISK_DEGREE'])
    fig7=plot7.get_figure()
    fig7.savefig(url+'/'+u'风险度')
    
    
    
    #交易量
    plot8=tt.plot.bar(x=['date'],y=['REAL_BUSINESS_AMOUNT'])
    ticklabels = ['']*len(tt.date)
    ticklabels[::21] = [item.strftime('%b %d\n%Y') for item in tt.date[::21]]
    plot8.xaxis.set_major_formatter(ticker.FixedFormatter(ticklabels))
    plt.gcf().autofmt_xdate()
    plt.show()
    fig8=plot8.get_figure()
    fig8.savefig(url+'/'+u'交易量')
    #持仓量
    plot9=tt.plot.bar(x=['date'],y=['HOLD_AMOUNT'])
    ticklabels = ['']*len(tt.date)
    ticklabels[::21] = [item.strftime('%b %d\n%Y') for item in tt.date[::21]]
    plot9.xaxis.set_major_formatter(ticker.FixedFormatter(ticklabels))
    plt.gcf().autofmt_xdate()
    plt.show()
    fig9=plot9.get_figure()
    fig9.savefig(url+'/'+u'持仓量')    
    
    
    #保证金占用
    plot10=tt.plot(x=['date'],y=['MARGIN_BALANCE'])
    fig10=plot10.get_figure()
    fig10.savefig(url+'/'+u'占用保证金')
    
    #输出客户数据
    CT.loc[acct].to_csv(url+'/'+str(acct)+'.csv')
 
def createDocx(acct,CT):
    #使用docx包，批量生成word文档
    document = Document()
    url=dir_url+str(acct).decode('UTF-8')
    print url    
    #标题    
    file_name= u'客户风险报告——客户号'+str(acct).decode('UTF-8')
    print file_name    
    document.add_heading(file_name, 1)  
    
    docRstyle1(document.paragraphs[0].runs[0])
       
    #表格，基础数据
    p1=document.add_heading(u'客户基础指标',2)
    docRstyle1(p1.runs[0])
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text =u'业绩指标'
    hdr_cells[1].text = u'数值'
    
    row_cells = table.add_row().cells
    row_cells[0].text =u'净值'
    row_cells[1].text =str(CT.loc[acct].last_nv)
    print 'style1'
    row_cells = table.add_row().cells
    row_cells[0].text =u'盈亏金额(元)'
    row_cells[1].text =str(CT.loc[acct].total_profit)
    print 'style2'
    row_cells = table.add_row().cells
    row_cells[0].text =u'交易量(张)'
    row_cells[1].text =str(CT.loc[acct].OPEN_AMOUNT)
    print 'style3'
    row_cells = table.add_row().cells
    row_cells[0].text =u'夏普率'
    row_cells[1].text =str(CT.loc[acct].sharp_ratio)
   
    row_cells = table.add_row().cells
    row_cells[0].text =u'初始交易时间'
    row_cells[1].text=str(CT.loc[acct].first_trade_day.date())
    
    row_cells = table.add_row().cells
    row_cells[0].text =u'最后交易时间'
    row_cells[1].text=str(CT.loc[acct].last_trade_day.date())    
    
    
    #图表，净值及盈亏
    p2=document.add_heading(u'图表-净值及盈亏',2)
    docRstyle1(p2.runs[0])
    document.add_picture(url+u'/净值.png', width=Inches(5))
    print 'style4'
    #document.add_picture(url+'净值.png', width=Inches(5))
   #图表，交易情况
    p3=document.add_heading(u'图表-交易情况',2)
    docRstyle1(p3.runs[0])
    document.add_picture(url+u'/交易量.png', width=Inches(5))  
    document.add_picture(url+u'/持仓量.png', width=Inches(5)) 
    print 'style35'

    #图表，希腊字母
    p4=document.add_heading(u'图表-希腊字母',2)
    docRstyle1(p4.runs[0])
    document.add_picture(url+u'/DELTA.png', width=Inches(5))  
    document.add_picture(url+u'/GAMMA.png', width=Inches(5))  
    document.add_picture(url+u'/VEGA.png', width=Inches(5))  
    document.add_picture(url+u'/THETA.png', width=Inches(5))      
    print 'style36'

    #图表，风险情况
    p5=document.add_heading(u'图表-交易情况',2)
    docRstyle1(p5.runs[0])
    document.add_picture(url+u'/风险度.png', width=Inches(5))  
    document.add_picture(url+u'/占用保证金.png', width=Inches(5))     
    print 'style37'
    
    #存储         
    document.save('D:/temp/doc/'+file_name+'.docx')    
    return

def docRstyle1(r1):
    r1.font.name=u'宋体'
    r1.font.color.rgb=RGBColor(0x24,0x24,0x24)      
    r = r1._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    return
    

__version__= '0.1'

if __name__ == '__main__':
    print 'this is in main,CAmodule'
    url1=u'D:/Work/日常事务/20160627-客户总体数据处理/allacct20160624_1.csv'
    url2=u'D:/Work/日常事务/20160627-客户总体数据处理/tradedetail.csv'            
    cdf=CAmodule.readRaredata(url1)
    tdf=CAmodule.readRaredata(url2)
    CAmodule.filterData(tdf)
    CAmodule.dealwithDate(cdf,tdf)
    CT=CAmodule.calCustomtransaction(CT,tdf)
    CT=CAmodule.calCustomNV(CT,cdf,tdf)    
    ii=CT.OPEN_AMOUNT.order().dropna().tail(10).index.get_values()